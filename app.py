import io
import json
import os
import re
from datetime import datetime, timezone
from functools import wraps
from urllib.parse import quote_plus

import firebase_admin
from docx import Document
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from firebase_admin import auth as firebase_auth, credentials, firestore
from google.cloud.firestore_v1.base_query import FieldFilter
from flask import (
    Flask, abort, jsonify, render_template, request, redirect, url_for,
    flash, send_file, session,
)
from google import genai
from google.genai import types
from openai import OpenAI
from pydantic import BaseModel
from pypdf import PdfReader
from werkzeug.security import generate_password_hash, check_password_hash
from xhtml2pdf import pisa

import adminstore
import applykit
import prepbank
import jobsources
import jobstore
import prepsets

load_dotenv()

app = Flask(__name__)
# Reads from environment; falls back to a dev-only value so local runs don't crash.
app.secret_key = os.environ.get("SECRET_KEY", "dev-only-secret-key-change-me")

GEMINI_MODEL = "gemini-3.5-flash"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
genai_client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

# Text generation (CV, cover letter, plagiarism write-up) runs on OpenRouter's
# free tier instead of Gemini, so it isn't capped by Gemini's 20 req/day quota.
# Get a key at https://openrouter.ai/keys - no billing required for :free models.
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "nvidia/nemotron-3-super-120b-a12b:free")
openrouter_client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=OPENROUTER_API_KEY,
) if OPENROUTER_API_KEY else None

RESUME_TEMPLATES = {"modern", "classic", "minimal"}

# Firebase Authentication, Google provider.
#
# The web config below is public by design - it ships in the page source of
# every Firebase app and is not a credential. What actually protects the flow
# is that the browser only ever returns a signed ID token, which the server
# verifies with the service account already loaded above. Access is controlled
# in the Firebase console (authorised domains + enabled providers), not by
# keeping these values quiet.
#
# Unset the API key and the Google button simply is not rendered; email
# sign-in is unaffected.
FIREBASE_WEB_API_KEY = os.environ.get("FIREBASE_WEB_API_KEY", "").strip()
FIREBASE_PROJECT_ID = os.environ.get("FIREBASE_PROJECT_ID", "").strip()
FIREBASE_AUTH_DOMAIN = os.environ.get("FIREBASE_AUTH_DOMAIN", "").strip()


def _firebase_web_config():
    """
    Client config for the Firebase JS SDK.

    project id and auth domain are derived from the service account when they
    are not set explicitly, so the only value that has to be supplied by hand
    is the web API key.
    """
    if not FIREBASE_WEB_API_KEY:
        return None

    project_id = FIREBASE_PROJECT_ID
    if not project_id:
        try:
            project_id = firebase_admin.get_app().project_id or ""
        except Exception:
            project_id = ""

    return {
        "apiKey": FIREBASE_WEB_API_KEY,
        "authDomain": FIREBASE_AUTH_DOMAIN or (f"{project_id}.firebaseapp.com" if project_id else ""),
        "projectId": project_id,
    }


def _openrouter_generate(prompt, json_mode=False):
    """Runs a single-turn text generation on OpenRouter. Raises if the key
    isn't configured or the API call fails - callers handle/report that."""
    kwargs = {}
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    completion = openrouter_client.chat.completions.create(
        model=OPENROUTER_MODEL,
        messages=[{"role": "user", "content": prompt}],
        **kwargs,
    )
    return completion.choices[0].message.content


def _init_firestore():
    """
    Loads Firebase Admin SDK credentials from either a JSON blob in
    FIREBASE_CREDENTIALS_JSON (used on Render, set as a secret env var) or a
    service account key file on disk (used locally). Returns None instead of
    raising if neither is configured yet, so the rest of the app can still
    run and surface a clear error only when a DB-backed route is hit.
    """
    try:
        if not firebase_admin._apps:
            cred_json = os.environ.get("FIREBASE_CREDENTIALS_JSON")
            if cred_json:
                cred = credentials.Certificate(json.loads(cred_json))
            else:
                cred_path = os.environ.get("FIREBASE_CREDENTIALS_PATH", "serviceAccountKey.json")
                cred = credentials.Certificate(cred_path)
            firebase_admin.initialize_app(cred)
        return firestore.client()
    except Exception as e:
        print(f"[WARN] Firebase/Firestore not configured: {e}")
        return None


db = _init_firestore()
adminstore.ensure_admin(db)


@app.template_filter('admin_date')
def _admin_date(epoch):
    """Epoch float -> readable date. Blank for the 0 that means 'never recorded'."""
    if not epoch:
        return '—'
    return datetime.fromtimestamp(epoch, tz=timezone.utc).strftime('%d %b %Y, %H:%M')


@app.context_processor
def inject_site_context():
    """
    Site name, footer signature and pending announcements, on every page.

    All three are served from adminstore's in-process cache, so this adds no
    Firestore reads to a normal page load.
    """
    settings = adminstore.get_settings(db)
    unseen = []
    if session.get('user_id'):
        unseen = adminstore.unseen_announcements(db, session.get('ann_seen_at', 0))
    return {
        'site': settings,
        'pending_announcements': unseen,
    }


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get('user_id'):
            # The apply endpoints are called with fetch(), so send them a JSON
            # 401 with somewhere to go - an HTML redirect would arrive as a
            # login page parsed as an API response.
            if request.is_json or request.path.startswith('/jobs/apply'):
                return jsonify({
                    "error": "Please log in to apply.",
                    "login_url": url_for('login', next=url_for('jobs')),
                }), 401
            flash('Please log in to continue.', 'error')
            return redirect(url_for('login', next=request.full_path))
        return view(*args, **kwargs)
    return wrapped


def admin_required(view):
    """
    Guards every /admin page except the login screen.

    Admin state lives under its own session key, so signing in as an admin does
    not grant a user session and vice versa - one is not an escalation of the
    other.
    """
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get('admin_email'):
            return redirect(url_for('admin_login', next=request.path))
        return view(*args, **kwargs)
    return wrapped


# --- ROUTES ---

@app.route('/')
def home():
    user_id = session.get('user_id')
    all_jobs = jobstore.all_jobs()

    # Static data, so the prep dashboard renders even on the cold-cache path.
    prep = {
        'prep_companies': _prep_companies(),
        'prep_sectors': _prep_sectors(),
        'prep_totals': _prep_totals(),
        # Admin-published openings, visible to everyone signed in or not.
        'posted_jobs': adminstore.list_job_posts(db, limit=6),
    }

    # Never block the landing page on a cold fetch (~8s across four boards) -
    # warm it in the background and let the section fill in on the next load.
    if not all_jobs:
        jobstore.warm_async(_fetch_live_jobs)
        return render_template(
            'index.html', preview_jobs=[], total_jobs=0,
            jobs_warming=jobstore.is_warming(), **prep,
        )

    profile = _user_job_profile(_load_profile_details())
    personalised = bool(user_id and profile['skills'])
    if personalised:
        jobsources.score_jobs(all_jobs, profile)
    applied_ids = jobstore.applied_job_ids(db, user_id) if user_id else set()

    return render_template(
        'index.html',
        preview_jobs=jobstore.preview_jobs(all_jobs, applied_ids, personalised),
        total_jobs=len(all_jobs),
        jobs_warming=False,
        personalised=personalised,
        **prep,
    )


def _prep_companies():
    """
    Every company that has questions, not just the ones prepsets knows about.

    A company created by adding a question in the admin panel had a working
    practice page but no card, so searching for it on /preparation found
    nothing. Synthesising a card here means anything with questions is
    listed - and therefore searchable.
    """
    cards = prepsets.list_companies()
    known = {c['slug'] for c in cards}

    for slug, name in prepbank.company_names(db).items():
        if slug in known:
            continue
        counts = prepbank.counts_for(db, slug)
        types = [t for t, n in counts['by_type'].items() if n]
        cards.append({
            "slug": slug,
            "name": name,
            "sector": "Added by admin",
            "difficulty": "Moderate",
            "focus": types,
            "rounds": [],
            "blurb": f"{counts['total']} practice question"
                     f"{'' if counts['total'] == 1 else 's'} added from the admin panel.",
            "initials": name[:2].upper(),
            "question_count": counts['total'],
            "round_count": 0,
        })

    cards.sort(key=lambda c: c['name'].lower())
    return cards


def _prep_sectors():
    return sorted({c['sector'] for c in _prep_companies()})


def _prep_totals():
    """
    Headline counts for the prep pages.

    prepsets only knows about its own interview Q&A; the practice bank holds
    an order of magnitude more. Counting one and labelling it "practice
    questions" understated the number by ~200, so the two are summed here.
    """
    totals = dict(prepsets.totals())
    totals['interview_questions'] = totals['questions']
    totals['practice_questions'] = prepbank.total_questions(db)
    totals['questions'] = totals['interview_questions'] + totals['practice_questions']
    return totals


@app.route('/preparation')
def preparation():
    """Standalone prep dashboard. Public, same as the section on the home page."""
    return render_template(
        'preparation.html',
        prep_companies=_prep_companies(),
        prep_sectors=_prep_sectors(),
        prep_totals=_prep_totals(),
    )


@app.route('/preparation/<slug>')
@login_required
def preparation_company(slug):
    """
    The practice desk for one company. Signed-in only - a guest clicking the
    same card gets the preview modal instead, which is why the card markup
    branches on the session rather than this route redirecting.
    """
    base = prepsets.get_set(slug, unlocked=True)
    names = prepbank.company_names(db)
    if base is None and slug not in names:
        abort(404)

    if base is None:
        # A company that exists only because the admin added questions for it.
        base = {
            "slug": slug, "name": names[slug], "sector": "Added by admin",
            "difficulty": "Moderate", "focus": [], "rounds": [],
            "blurb": "", "initials": names[slug][:2].upper(),
        }

    # Interview Q&A and practice problems are different shapes; flatten them
    # into one list so the left rail is a single scrollable index.
    items = []
    for q in prepbank.questions_for(db, slug):
        items.append({
            "id": q.get('id'),
            "kind": q.get('source', 'bank'),
            "type": q.get('type', 'Coding'),
            "title": q.get('title') or (q.get('prompt') or '')[:60],
            "difficulty": q.get('difficulty', 'Medium'),
            "prompt": q.get('prompt', ''),
            "url": q.get('url', ''),
            "example": q.get('example') or {},
            "options": q.get('options') or [],
            "answer": q.get('answer', ''),
            "explanation": q.get('explanation', ''),
            "example_function": q.get('example_function', ''),
            "acceptance_pct": q.get('acceptance_pct'),
            "frequency_pct": q.get('frequency_pct'),
        })

    for index, q in enumerate(base.get('questions') or []):
        items.append({
            "id": f"interview:{slug}:{index}",
            "kind": "interview",
            "type": q.get('type', 'Interview'),
            "title": q.get('prompt', '')[:70],
            "difficulty": "Interview",
            "prompt": q.get('prompt', ''),
            "url": "",
            "example": {},
            "options": [],
            "answer": q.get('answer', ''),
            "explanation": "",
            "example_function": "",
            "acceptance_pct": None,
            "frequency_pct": None,
        })

    return render_template(
        'practice.html',
        company=base,
        items=items,
        counts=prepbank.counts_for(db, slug),
    )


@app.route('/prep/<slug>/set')
def prep_set(slug):
    """
    One company's practice set.

    Deliberately not behind @login_required: a guest gets a real response, just
    a shortened one. The withheld questions are dropped in prepsets.get_set(),
    so the lock survives someone deleting the blur in devtools.
    """
    payload = prepsets.get_set(slug, unlocked=bool(session.get('user_id')))
    if payload is None:
        return jsonify({"error": "Unknown company."}), 404

    if not payload['unlocked']:
        payload['login_url'] = url_for('login', next=url_for('home') + '#prep')
    return jsonify(payload)


def _render_auth(mode):
    """Both /login and /register render the same page; mode picks the panel."""
    return render_template(
        'auth.html',
        mode=mode,
        next=request.args.get('next', ''),
        firebase_config=_firebase_web_config(),
    )


def _safe_next(target):
    """
    Only ever redirect within this site.

    `next` arrives from a query string or a JSON body, so without this an
    attacker could hand someone a login link that bounces to their own domain
    after a real sign-in. A leading `//` is excluded too - browsers read that
    as protocol-relative and would leave the site.
    """
    if not target or not target.startswith('/') or target.startswith('//'):
        return None
    return target


def _start_session(email, user):
    session['user_id'] = email
    session['username'] = user.get('username') or email.split('@')[0]
    # Carried in the session so the announcement bell costs no Firestore read
    # per page; persisted on the user when dismissed.
    session['ann_seen_at'] = user.get('ann_seen_at', 0)


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        if db is None:
            flash('Database is not configured on the server.', 'error')
            return _render_auth('register')

        username = request.form['username']
        email = request.form['email'].strip().lower()
        password = request.form['password']

        user_ref = db.collection('users').document(email)
        if user_ref.get().exists:
            flash('Email already exists.', 'error')
        else:
            user_ref.set({
                'username': username,
                'email': email,
                'password': generate_password_hash(password),
                'auth_provider': 'password',
                'created_at': firestore.SERVER_TIMESTAMP,
            })
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))

    return _render_auth('register')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email'].strip().lower()
        password = request.form['password']

        user = None
        if db is not None:
            doc = db.collection('users').document(email).get()
            if doc.exists:
                user = doc.to_dict()

        # Google-created accounts have no password hash, so check for one
        # before comparing - and say why, rather than "invalid password" on an
        # account that never had one.
        if user and not user.get('password'):
            flash('That account was created with Google — use the Google button above.', 'error')
        elif user and check_password_hash(user['password'], password):
            _start_session(email, user)
            flash('Login successful!', 'success')
            return redirect(_safe_next(request.form.get('next')) or url_for('home'))
        else:
            flash('Invalid email or password.', 'error')

    return _render_auth('login')


@app.route('/auth/google', methods=['POST'])
def auth_google():
    """
    Verifies a Firebase ID token and signs the user in, registering them on
    first use.

    The browser does the Google popup and hands back a token; this checks the
    signature and audience against our own Firebase project using the service
    account. Nothing here trusts the email the page claims - it is read out of
    the verified token.
    """
    if not FIREBASE_WEB_API_KEY:
        return jsonify({"error": "Google sign-in is not configured on this server."}), 503
    if db is None:
        return jsonify({"error": "Database is not configured on the server."}), 503

    credential = (request.get_json(silent=True) or {}).get('credential')
    if not credential:
        return jsonify({"error": "Missing sign-in token."}), 400

    try:
        claims = firebase_auth.verify_id_token(credential)
    except Exception as e:
        return jsonify({"error": f"Could not verify that Google account: {e}"}), 401

    if not claims.get('email'):
        return jsonify({"error": "That Google account exposed no email address."}), 401
    if not claims.get('email_verified'):
        return jsonify({"error": "That Google account has no verified email."}), 401

    email = claims['email'].strip().lower()
    user_ref = db.collection('users').document(email)
    doc = user_ref.get()

    if doc.exists:
        user = doc.to_dict()
    else:
        # No password field: this account can only ever be reached through
        # Google, and /login says so rather than failing as a bad password.
        user = {
            'username': claims.get('name') or email.split('@')[0],
            'email': email,
            'auth_provider': 'google',
            'created_at': firestore.SERVER_TIMESTAMP,
        }
        user_ref.set(user)

    _start_session(email, user)
    target = _safe_next((request.get_json(silent=True) or {}).get('next'))
    return jsonify({"redirect": target or url_for('home')})


@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('home'))


def _user_job_profile(details):
    """Turns the saved profile details into the shape the matcher and the
    apply prefill both expect."""
    skills = [s.strip() for s in re.split(r"[,\n;]+", details.get('skills') or '') if s.strip()]
    titles = [t.strip() for t in re.split(r"[,\n;]+", details.get('desired_titles') or '') if t.strip()]
    locations = [details.get('location') or '', 'remote']
    return {
        'full_name': details.get('full_name') or session.get('username', ''),
        'email': session.get('user_id', ''),
        'phone': details.get('phone') or '',
        'location': details.get('location') or '',
        'linkedin': details.get('linkedin') or '',
        'github': details.get('github') or '',
        'portfolio': details.get('other') or '',
        'skills': skills,
        'desired_titles': titles,
        'desired_locations': [loc for loc in locations if loc],
    }


def _load_profile_details():
    if db is None or not session.get('user_id'):
        return {}
    doc = db.collection('users').document(session['user_id']).get()
    return (doc.to_dict().get('details') or {}) if doc.exists else {}


def _fetch_live_jobs():
    return jobsources.fetch_all(jobsources.load_config())


@app.route('/jobs')
def jobs():
    """Public: anyone can browse and filter listings. Applying needs an account,
    which the template asks for at the point of use."""
    user_id = session.get('user_id')
    profile = _user_job_profile(_load_profile_details())

    all_jobs = jobstore.all_jobs()
    state = jobstore.cache_state()

    # First visit after a restart has an empty cache - pull one set so the page
    # is never blank, rather than making the user guess they must click Fetch.
    if not all_jobs:
        try:
            fetched, errors = _fetch_live_jobs()
            jobstore.store_jobs(fetched, errors)
            all_jobs = jobstore.all_jobs()
            state = jobstore.cache_state()
        except Exception as exc:  # noqa: BLE001
            state['errors'] = [str(exc)]

    jobsources.score_jobs(all_jobs, profile)
    applied_ids = jobstore.applied_job_ids(db, user_id) if user_id else set()

    results = jobstore.filter_jobs(
        all_jobs,
        search=request.args.get('q', ''),
        source=request.args.get('source', ''),
        category=request.args.get('category', ''),
        applied_ids=applied_ids,
        hide_applied=request.args.get('hide_applied') == '1',
    )

    # No runtime capability check any more - the apply kit is pure Python, so
    # it works on every host rather than only where a browser is installed.
    return render_template(
        'jobs.html',
        jobs=results,
        total_cached=state['count'],
        fetched_at=state['fetched_at'],
        is_stale=state['is_stale'],
        fetch_errors=state['errors'],
        sources=jobstore.distinct(all_jobs, 'source'),
        categories=jobstore.distinct(all_jobs, 'category'),
        search=request.args.get('q', ''),
        selected_source=request.args.get('source', ''),
        selected_category=request.args.get('category', ''),
        hide_applied=request.args.get('hide_applied') == '1',
        applied_count=len(applied_ids),
        has_skills=bool(profile['skills']),
        direct_apply_ats=sorted(applykit.SUPPORTED_ATS),
        # Offered as the CV to attach when applying; most recent is the default.
        saved_resumes=_fetch_user_docs('resumes') if user_id else [],
        signed_in=bool(user_id),
    )


@app.route('/jobs/fetch', methods=['POST'])
def jobs_fetch():
    try:
        fetched, errors = _fetch_live_jobs()
    except Exception as exc:  # noqa: BLE001
        flash(f'Could not refresh listings: {exc}', 'error')
        return redirect(url_for('jobs'))

    jobstore.store_jobs(fetched, errors)
    if errors:
        flash(f"Fetched {len(fetched)} jobs, but some boards failed: {'; '.join(errors)}", 'error')
    else:
        flash(f'Fetched {len(fetched)} live job postings.', 'success')
    return redirect(url_for('jobs', **{k: v for k, v in request.args.items()}))


def _resume_pdf_for_apply(resume_id):
    """Renders one of the user's saved resumes to PDF bytes so the apply
    service can upload it. Falls back to their most recent resume when the
    caller didn't pick one, so applications aren't sent without a CV attached.
    Returns (bytes, filename) or (None, '')."""
    if db is None:
        return None, ''

    row = None
    if resume_id:
        doc = db.collection('resumes').document(resume_id).get()
        if doc.exists and doc.to_dict().get('user_id') == session['user_id']:
            row = doc.to_dict()
    if row is None:
        saved = _fetch_user_docs('resumes')
        if not saved:
            return None, ''
        row = saved[0]
    try:
        html = _render_resume_html(row.get('template', 'modern'), row.get('generated_content') or {})
        return _html_to_pdf_bytes(html), f"{_slugify(row.get('full_name') or 'resume')}.pdf"
    except Exception:  # noqa: BLE001 - apply can still proceed without a CV attached
        return None, ''


@app.route('/jobs/<job_id>/apply-kit')
@login_required
def jobs_apply_kit(job_id):
    """Everything needed to finish one application: the (pre-filled where the
    board supports it) form URL, the values to enter, and whether a tailored CV
    is available to download."""
    job = jobstore.get_job(job_id)
    if job is None:
        return jsonify({"error": "That posting is no longer in the current listing set. Refresh and try again."}), 404

    profile = _user_job_profile(_load_profile_details())
    if not profile['full_name']:
        return jsonify({"error": "Add your name to your profile before applying."}), 400

    resume_bytes, _ = _resume_pdf_for_apply(request.args.get('resume_id'))
    kit = applykit.build_kit(job, profile, has_resume=bool(resume_bytes))
    kit['resume_url'] = url_for('jobs_apply_cv', job_id=job_id,
                                resume_id=request.args.get('resume_id', ''))
    return jsonify(kit)


@app.route('/jobs/<job_id>/apply-cv.pdf')
@login_required
def jobs_apply_cv(job_id):
    """The tailored CV to attach, rendered on demand from a saved resume."""
    job = jobstore.get_job(job_id)
    resume_bytes, filename = _resume_pdf_for_apply(request.args.get('resume_id'))
    if not resume_bytes:
        abort(404)
    company = _slugify((job or {}).get('company') or 'application')
    return send_file(
        io.BytesIO(resume_bytes), mimetype='application/pdf',
        as_attachment=True, download_name=f"{_slugify(filename.rsplit('.', 1)[0])}-{company}.pdf",
    )



@app.route('/jobs/mark-applied', methods=['POST'])
@login_required
def jobs_mark_applied():
    """Records an application. 'submitted' comes from the apply kit, where the
    user confirmed they sent it after we prepared it; 'manual' is someone
    ticking off a job they found and applied to on their own."""
    data = request.get_json(silent=True) or {}
    job = jobstore.get_job((data.get('job_id') or '').strip())
    if job is None:
        return jsonify({"error": "Unknown posting."}), 404

    status = 'submitted' if data.get('status') == 'submitted' else 'manual'
    try:
        saved = jobstore.record_application(
            db, session['user_id'], job, status=status,
            match_percentage=job.get('match_percentage', 0),
        )
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 503
    return jsonify({"ok": True, "application": {"id": saved["id"], "status": saved["status"]}})


@app.route('/applications/<application_id>/delete', methods=['POST'])
@login_required
def application_delete(application_id):
    if not jobstore.delete_application(db, session['user_id'], application_id):
        abort(404)
    flash('Application removed from your tracker.', 'success')
    return redirect(url_for('profile'))


COURSE_PROVIDERS = [
    {"name": "Coursera", "icon": "🎓", "url": "https://www.coursera.org/search?query={q}"},
    {"name": "Udemy", "icon": "💻", "url": "https://www.udemy.com/courses/search/?q={q}"},
    {"name": "edX", "icon": "📘", "url": "https://www.edx.org/search?q={q}"},
    {"name": "YouTube", "icon": "▶️", "url": "https://www.youtube.com/results?search_query={q}"},
]


def _course_links_for_keyword(keyword):
    links = []
    for provider in COURSE_PROVIDERS:
        query = f"{keyword} course" if provider["name"] == "YouTube" else keyword
        links.append({
            "name": provider["name"],
            "icon": provider["icon"],
            "url": provider["url"].format(q=quote_plus(query)),
        })
    return links


@app.route('/courses')
def courses():
    raw = request.args.get('keywords', '')
    seen = set()
    keywords = []
    for part in raw.split(','):
        keyword = part.strip()
        if keyword and keyword.lower() not in seen:
            seen.add(keyword.lower())
            keywords.append(keyword)
    keywords = keywords[:20]

    keyword_courses = [
        {"keyword": keyword, "links": _course_links_for_keyword(keyword)}
        for keyword in keywords
    ]
    return render_template('courses.html', keyword_courses=keyword_courses, raw_keywords=raw)


def _extract_json_array(text):
    match = re.search(r"\[.*\]", text, re.DOTALL)
    if not match:
        raise ValueError("Model response did not contain a JSON array.")
    return json.loads(match.group(0))


def _extract_json_object(text):
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("Model response did not contain a JSON object.")
    return json.loads(match.group(0))




@app.route('/jobs/generate-cv', methods=['POST'])
def jobs_generate_cv():
    if not openrouter_client:
        return jsonify({"error": "OPENROUTER_API_KEY is not configured on the server."}), 503

    data = request.get_json(silent=True) or {}
    job_title = (data.get('title') or '').strip()
    company = (data.get('company') or '').strip()
    job_description = (data.get('job_description') or '').strip()
    keywords = data.get('keywords') or []
    skills = (data.get('skills') or '').strip()

    if not job_description:
        return jsonify({"error": "Job description is required."}), 400

    prompt = f"""Create a tailored, ATS-friendly CV/resume in plain text for a candidate
applying to this specific job.

Job Title: {job_title}
Company: {company or "N/A"}
Job Description:
{job_description}

Keywords to naturally incorporate: {", ".join(keywords) if keywords else "N/A"}

Candidate's current skills / experience / background:
{skills or "Not provided - infer a reasonable background suitable for this role."}

Write a complete CV with: a professional summary tailored to this role, a
skills section prominently featuring the keywords above, and a suggested
work experience section with bullet points. Do not fabricate specific
employers or degrees; use placeholders like [Your Previous Company] where
actual history is unknown. Output plain text only, no markdown symbols."""

    try:
        cv_text = _openrouter_generate(prompt).strip()
    except Exception as e:
        return jsonify({"error": f"Failed to generate CV: {e}"}), 502

    return jsonify({"cv": cv_text})


# --- GENERATE CV (BUILDER, TEMPLATES, EXPORTS, PROFILE) ---

class ResumeContact(BaseModel):
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    portfolio: str = ""


class ResumeExperience(BaseModel):
    title: str = ""
    company: str = ""
    dates: str = ""
    bullets: list[str] = []


class ResumeEducation(BaseModel):
    degree: str = ""
    school: str = ""
    dates: str = ""


class ResumeProject(BaseModel):
    name: str = ""
    description: str = ""


class ResumeContent(BaseModel):
    full_name: str = ""
    target_role: str = ""
    contact: ResumeContact = ResumeContact()
    summary: str = ""
    skills: list[str] = []
    experience: list[ResumeExperience] = []
    education: list[ResumeEducation] = []
    certifications: list[str] = []
    projects: list[ResumeProject] = []


class ResumeSectionChange(BaseModel):
    section: str = ""
    change: str = ""


class ResumeAnalysis(BaseModel):
    before_score: int = 0
    before_reasons: list[str] = []
    after_score: int = 0
    after_reasons: list[str] = []
    keywords_added: list[str] = []
    section_changes: list[ResumeSectionChange] = []
    action_items: list[str] = []
    optimized_resume_markdown: str = ""


def _slugify(text):
    slug = re.sub(r'[^a-zA-Z0-9]+', '-', text or '').strip('-').lower()
    return slug or 'resume'


def _generate_resume_content(input_data):
    prompt = f"""You are an expert resume writer. Build an ATS-friendly resume from the
candidate's raw input below. Rewrite experience bullets to be concise,
action-oriented, quantified where reasonable, and keyword-rich for the
target role. Do not fabricate employers, dates, or degrees beyond what is
given here - only improve the wording and structure of what's provided. If
the raw input already contains bracketed placeholders (e.g.
"[Your Previous Company]", "[Graduation Year]"), keep those placeholders in
the output literally instead of dropping them - a section left as
placeholders is more useful to the candidate than an empty section.

Candidate input:
Full Name: {input_data.get('full_name', '')}
Target Role: {input_data.get('target_role', '')}
Email: {input_data.get('email', '')}
Phone: {input_data.get('phone', '')}
Location: {input_data.get('location', '')}
LinkedIn: {input_data.get('linkedin', '')}
Portfolio: {input_data.get('portfolio', '')}
Professional Summary (raw, optional): {input_data.get('summary', '')}
Skills (raw): {input_data.get('skills', '')}
Work Experience (raw): {input_data.get('experience', '')}
Education (raw): {input_data.get('education', '')}
Certifications (raw, optional): {input_data.get('certifications', '')}
Projects (raw, optional): {input_data.get('projects', '')}

Use empty strings/arrays for fields with no data. Do not invent employers,
dates, or degrees beyond what is given.

Respond with ONLY a JSON object (no markdown, no commentary) with exactly
these keys:
- "full_name": string
- "target_role": string
- "contact": object with "email", "phone", "location", "linkedin", "portfolio" (strings)
- "summary": string
- "skills": array of strings
- "experience": array of objects with "title", "company", "dates" (strings) and "bullets" (array of strings)
- "education": array of objects with "degree", "school", "dates" (strings)
- "certifications": array of strings
- "projects": array of objects with "name", "description" (strings)"""

    raw = _openrouter_generate(prompt, json_mode=True)
    data = _extract_json_object(raw)
    return ResumeContent(**data).model_dump()


MAX_ANALYZE_RESUME_CHARS = 12000
MAX_ANALYZE_JD_CHARS = 6000


def _analyze_and_tailor_resume(resume_text, job_description):
    resume_text = resume_text[:MAX_ANALYZE_RESUME_CHARS]
    job_description = job_description[:MAX_ANALYZE_JD_CHARS]

    prompt = f"""You are an expert ATS (Applicant Tracking System) resume consultant.
A candidate has given you their current resume and a target job description. Do all
of the following:

1. Score the ORIGINAL resume as-is against the job description for ATS compatibility
   and keyword/skill match, as an integer 0-100 ("before_score"), and give 3-5 short,
   specific reasons for that score ("before_reasons") - e.g. missing keywords, weak
   quantification, poor formatting for parsing, irrelevant content.
2. Rewrite the resume into an ATS-optimized version tailored to this job: reorganize
   and rephrase for clarity and keyword density, surface relevant skills/experience
   the candidate already has that map to the job description, and tighten bullet
   points to be action-oriented and quantified where the original already implies a
   metric. Do NOT fabricate employers, job titles, dates, degrees, or achievements
   that aren't supported by the original resume - if a requirement in the job
   description is genuinely unmet, leave a bracketed placeholder like
   "[Add specific metric here]" instead of inventing one.
3. Score the OPTIMIZED resume the same way as step 1, as "after_score" (integer
   0-100) with "after_reasons" explaining the improvement.
4. List the specific keywords/skills from the job description that were newly
   incorporated into the optimized resume and were missing or weak in the original
   ("keywords_added").
5. Summarize what changed section-by-section ("section_changes") - at minimum cover
   Summary, Experience, and Skills if the resume has them, plus any other section that
   changed meaningfully. Each entry has a "section" name and a one-to-two sentence
   "change" description of what was revised and why.
6. List manual action items for the candidate - things you could not safely do for
   them ("action_items"): missing certifications/requirements from the job
   description, places that need real numbers/metrics only the candidate knows, gaps
   that need clarification, etc. Empty array if none.

Candidate's original resume:
\"\"\"
{resume_text}
\"\"\"

Target job description:
\"\"\"
{job_description}
\"\"\"

Respond with ONLY a JSON object (no markdown, no commentary) with exactly these keys:
- "before_score": integer 0-100
- "before_reasons": array of strings
- "after_score": integer 0-100
- "after_reasons": array of strings
- "keywords_added": array of strings
- "section_changes": array of objects with "section" and "change" (strings)
- "action_items": array of strings
- "optimized_resume_markdown": string - the full rewritten resume formatted as clean
  Markdown (headings, bullet lists), ready to copy-paste"""

    raw = _openrouter_generate(prompt, json_mode=True)
    data = _extract_json_object(raw)
    return ResumeAnalysis(**data).model_dump()


def _render_resume_html(template, resume):
    if template not in RESUME_TEMPLATES:
        template = 'modern'
    return render_template(f'resume_templates/{template}.html', resume=resume)


def _html_to_pdf_bytes(html):
    buf = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(html), dest=buf)
    if result.err:
        raise RuntimeError("Failed to render PDF.")
    return buf.getvalue()


_TEMPLATE_ACCENTS = {
    "modern": RGBColor(0x1F, 0x3A, 0x5F),
    "classic": RGBColor(0x00, 0x00, 0x00),
    "minimal": RGBColor(0x33, 0x33, 0x33),
}


def _build_resume_docx(template, resume):
    doc = Document()
    accent = _TEMPLATE_ACCENTS.get(template, _TEMPLATE_ACCENTS["modern"])
    font_name = 'Georgia' if template == 'classic' else 'Calibri'

    base_style = doc.styles['Normal']
    base_style.font.name = font_name
    base_style.font.size = Pt(10.5)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.get('full_name', ''))
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = accent

    if resume.get('target_role'):
        role_run = doc.add_paragraph().add_run(resume['target_role'])
        role_run.font.size = Pt(12)
        role_run.italic = True

    contact = resume.get('contact') or {}
    contact_line = ' | '.join(filter(None, [
        contact.get('email'), contact.get('phone'), contact.get('location'),
        contact.get('linkedin'), contact.get('portfolio'),
    ]))
    if contact_line:
        doc.add_paragraph().add_run(contact_line).font.size = Pt(9.5)

    def add_heading(text):
        h = doc.add_paragraph()
        run = h.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(2)

    if resume.get('summary'):
        add_heading('Summary')
        doc.add_paragraph(resume['summary'])

    if resume.get('skills'):
        add_heading('Skills')
        doc.add_paragraph(' • '.join(resume['skills']))

    if resume.get('experience'):
        add_heading('Experience')
        for job in resume['experience']:
            jp = doc.add_paragraph()
            jp.add_run(f"{job.get('title', '')} — {job.get('company', '')}").bold = True
            if job.get('dates'):
                jp.add_run(f"   ({job['dates']})").italic = True
            for bullet in job.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet')

    if resume.get('education'):
        add_heading('Education')
        for edu in resume['education']:
            ep = doc.add_paragraph()
            ep.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}").bold = True
            if edu.get('dates'):
                ep.add_run(f"   ({edu['dates']})").italic = True

    if resume.get('certifications'):
        add_heading('Certifications')
        for cert in resume['certifications']:
            doc.add_paragraph(cert, style='List Bullet')

    if resume.get('projects'):
        add_heading('Projects')
        for proj in resume['projects']:
            pp = doc.add_paragraph()
            pp.add_run(proj.get('name', '')).bold = True
            if proj.get('description'):
                doc.add_paragraph(proj['description'])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route('/generate-cv')
@login_required
def generate_cv():
    return render_template('generate_cv.html')


@app.route('/generate-cv/build', methods=['POST'])
@login_required
def generate_cv_build():
    if not openrouter_client:
        return jsonify({"error": "OPENROUTER_API_KEY is not configured on the server."}), 503

    input_data = request.get_json(silent=True) or {}
    if not (input_data.get('full_name') and input_data.get('target_role')):
        return jsonify({"error": "Full name and target role are required."}), 400

    try:
        resume = _generate_resume_content(input_data)
    except Exception as e:
        return jsonify({"error": f"Failed to generate resume: {e}"}), 502

    return jsonify({"resume": resume})


@app.route('/generate-cv/render', methods=['POST'])
@login_required
def generate_cv_render():
    data = request.get_json(silent=True) or {}
    return _render_resume_html(data.get('template', 'modern'), data.get('resume') or {})


@app.route('/generate-cv/download/pdf', methods=['POST'])
@login_required
def generate_cv_download_pdf():
    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    resume = data.get('resume') or {}
    try:
        pdf_bytes = _html_to_pdf_bytes(_render_resume_html(template, resume))
    except Exception as e:
        return jsonify({"error": f"Failed to build PDF: {e}"}), 500
    return send_file(
        io.BytesIO(pdf_bytes), mimetype='application/pdf',
        as_attachment=True, download_name=f"{_slugify(resume.get('full_name'))}.pdf",
    )


@app.route('/generate-cv/download/docx', methods=['POST'])
@login_required
def generate_cv_download_docx():
    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    resume = data.get('resume') or {}
    try:
        docx_bytes = _build_resume_docx(template, resume)
    except Exception as e:
        return jsonify({"error": f"Failed to build Word document: {e}"}), 500
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name=f"{_slugify(resume.get('full_name'))}.docx",
    )


@app.route('/generate-cv/save', methods=['POST'])
@login_required
def generate_cv_save():
    if db is None:
        return jsonify({"error": "Database is not configured on the server."}), 503

    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    input_data = data.get('input_data') or {}
    resume = data.get('resume') or {}

    if not resume:
        return jsonify({"error": "No generated resume to save."}), 400

    doc_ref = db.collection('resumes').document()
    doc_ref.set({
        'user_id': session['user_id'],
        'template': template,
        'full_name': resume.get('full_name', ''),
        'target_role': resume.get('target_role', ''),
        'input_data': input_data,
        'generated_content': resume,
        'created_at': firestore.SERVER_TIMESTAMP,
    })

    return jsonify({"id": doc_ref.id})


@app.route('/generate-cv/analyze', methods=['POST'])
@login_required
def generate_cv_analyze():
    if not openrouter_client:
        return jsonify({"error": "OPENROUTER_API_KEY is not configured on the server."}), 503

    if request.content_type and 'multipart/form-data' in request.content_type:
        job_description = (request.form.get('job_description') or '').strip()
        uploaded = request.files.get('resume_file')
        if uploaded and uploaded.filename:
            try:
                resume_text = _extract_text_from_upload(uploaded)
            except ValueError as e:
                return jsonify({"error": str(e)}), 400
        else:
            resume_text = request.form.get('resume_text') or ''
    else:
        data = request.get_json(silent=True) or {}
        resume_text = data.get('resume_text') or ''
        job_description = (data.get('job_description') or '').strip()

    resume_text = resume_text.strip()
    if not resume_text or not job_description:
        return jsonify({"error": "Please provide both your resume and the target job description."}), 400

    try:
        analysis = _analyze_and_tailor_resume(resume_text, job_description)
    except Exception as e:
        return jsonify({"error": f"Failed to analyze resume: {e}"}), 502

    return jsonify({"analysis": analysis})


@app.route('/generate-cv/analyze/save', methods=['POST'])
@login_required
def generate_cv_analyze_save():
    if db is None:
        return jsonify({"error": "Database is not configured on the server."}), 503

    data = request.get_json(silent=True) or {}
    job_description = data.get('job_description') or ''
    analysis = data.get('analysis') or {}

    if not analysis:
        return jsonify({"error": "No analysis to save."}), 400

    doc_ref = db.collection('resume_analyses').document()
    doc_ref.set({
        'user_id': session['user_id'],
        'job_description': job_description,
        'before_score': analysis.get('before_score', 0),
        'after_score': analysis.get('after_score', 0),
        'generated_content': analysis,
        'created_at': firestore.SERVER_TIMESTAMP,
    })

    return jsonify({"id": doc_ref.id})


@app.route('/analysis/<analysis_id>/delete', methods=['POST'])
@login_required
def analysis_delete(analysis_id):
    _get_owned_doc('resume_analyses', analysis_id)
    db.collection('resume_analyses').document(analysis_id).delete()
    return jsonify({"ok": True})


def _fetch_user_docs(collection):
    if db is None:
        return []
    docs = db.collection(collection).where(filter=FieldFilter('user_id', '==', session['user_id'])).stream()
    items = []
    for d in docs:
        data = d.to_dict()
        data['id'] = d.id
        items.append(data)
    items.sort(key=lambda r: r.get('created_at') or datetime.min.replace(tzinfo=timezone.utc), reverse=True)
    return items


def _get_owned_doc(collection, doc_id):
    if db is None:
        abort(404)
    doc = db.collection(collection).document(doc_id).get()
    if not doc.exists or doc.to_dict().get('user_id') != session['user_id']:
        abort(404)
    data = doc.to_dict()
    data['id'] = doc.id
    return data


@app.route('/profile')
@login_required
def profile():
    resumes = _fetch_user_docs('resumes')
    cover_letters = _fetch_user_docs('cover_letters')
    analyses = _fetch_user_docs('resume_analyses')
    analyses_markdown_map = {
        a['id']: (a.get('generated_content') or {}).get('optimized_resume_markdown', '')
        for a in analyses
    }
    profile_details = _load_profile_details()
    applications = jobstore.list_applications(db, session['user_id'])
    return render_template(
        'profile.html', resumes=resumes, cover_letters=cover_letters, analyses=analyses,
        analyses_markdown_map=analyses_markdown_map, profile_details=profile_details,
        applications=applications, analytics=jobstore.build_analytics(applications),
    )


@app.route('/profile/details/save', methods=['POST'])
@login_required
def profile_details_save():
    if db is None:
        return jsonify({"error": "Database is not configured on the server."}), 503

    data = request.get_json(silent=True) or {}
    details = {
        'full_name': (data.get('full_name') or '').strip(),
        'id_card_number': (data.get('id_card_number') or '').strip(),
        'college_name': (data.get('college_name') or '').strip(),
        'location': (data.get('location') or '').strip(),
        'github': (data.get('github') or '').strip(),
        'linkedin': (data.get('linkedin') or '').strip(),
        'other': (data.get('other') or '').strip(),
        'about_me': (data.get('about_me') or '').strip(),
        # Used by Find Jobs: phone pre-fills applications, skills and desired
        # titles drive the match percentage.
        'phone': (data.get('phone') or '').strip(),
        'skills': (data.get('skills') or '').strip(),
        'desired_titles': (data.get('desired_titles') or '').strip(),
    }
    db.collection('users').document(session['user_id']).update({'details': details})
    return jsonify({"ok": True, "details": details})


@app.route('/resume/<resume_id>/download/pdf')
@login_required
def resume_download_pdf(resume_id):
    row = _get_owned_doc('resumes', resume_id)
    pdf_bytes = _html_to_pdf_bytes(_render_resume_html(row['template'], row['generated_content']))
    return send_file(
        io.BytesIO(pdf_bytes), mimetype='application/pdf',
        as_attachment=True, download_name=f"{_slugify(row['full_name'])}.pdf",
    )


@app.route('/resume/<resume_id>/download/docx')
@login_required
def resume_download_docx(resume_id):
    row = _get_owned_doc('resumes', resume_id)
    docx_bytes = _build_resume_docx(row['template'], row['generated_content'])
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name=f"{_slugify(row['full_name'])}.docx",
    )


@app.route('/resume/<resume_id>/delete', methods=['POST'])
@login_required
def resume_delete(resume_id):
    _get_owned_doc('resumes', resume_id)
    db.collection('resumes').document(resume_id).delete()
    return jsonify({"ok": True})


# --- GENERATE COVER LETTER (BUILDER, TEMPLATES, EXPORTS, PROFILE) ---

COVER_LETTER_TEMPLATES = {"modern", "classic", "minimal"}


class CoverLetterContent(BaseModel):
    full_name: str = ""
    target_role: str = ""
    company_name: str = ""
    contact: ResumeContact = ResumeContact()
    salutation: str = ""
    body_paragraphs: list[str] = []
    closing: str = ""


def _generate_cover_letter_content(input_data):
    prompt = f"""You are an expert cover letter writer. Write a concise, ATS-friendly
cover letter from the candidate's raw input below, tailored to the target
role and company. Keep it to 3-4 short paragraphs: an opening hook, why
they're a fit (tied to the job description if given), a highlight or two
from their background, and a confident closing call to action. Do not
fabricate employers, achievements, or facts beyond what is given here -
only improve the wording and structure.

Candidate input:
Full Name: {input_data.get('full_name', '')}
Target Role: {input_data.get('target_role', '')}
Company Name: {input_data.get('company_name', '')}
Email: {input_data.get('email', '')}
Phone: {input_data.get('phone', '')}
Location: {input_data.get('location', '')}
LinkedIn: {input_data.get('linkedin', '')}
Portfolio: {input_data.get('portfolio', '')}
Hiring Manager Name (optional): {input_data.get('hiring_manager', '')}
Job Description / Key Requirements (raw, optional): {input_data.get('job_description', '')}
Your Background / Key Achievements (raw): {input_data.get('background', '')}

Use "Dear Hiring Manager," as the salutation if no hiring manager name is
given, otherwise "Dear {{name}},". body_paragraphs is a list of 3-4
plain-text paragraphs (no headers, no markdown). closing is a short
sign-off phrase like "Sincerely," - do not repeat the candidate's name in
it, that is rendered separately.

Respond with ONLY a JSON object (no markdown, no commentary) with exactly
these keys:
- "full_name": string
- "target_role": string
- "company_name": string
- "contact": object with "email", "phone", "location", "linkedin", "portfolio" (strings)
- "salutation": string
- "body_paragraphs": array of strings
- "closing": string"""

    raw = _openrouter_generate(prompt, json_mode=True)
    data = _extract_json_object(raw)
    return CoverLetterContent(**data).model_dump()


def _render_cover_letter_html(template, letter):
    if template not in COVER_LETTER_TEMPLATES:
        template = 'modern'
    return render_template(
        f'cover_letter_templates/{template}.html',
        letter=letter, today=datetime.now().strftime('%B %d, %Y'),
    )


def _build_cover_letter_docx(template, letter):
    doc = Document()
    accent = _TEMPLATE_ACCENTS.get(template, _TEMPLATE_ACCENTS["modern"])
    font_name = 'Georgia' if template == 'classic' else 'Calibri'

    base_style = doc.styles['Normal']
    base_style.font.name = font_name
    base_style.font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(letter.get('full_name', ''))
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = accent

    contact = letter.get('contact') or {}
    contact_line = ' | '.join(filter(None, [
        contact.get('email'), contact.get('phone'), contact.get('location'),
        contact.get('linkedin'), contact.get('portfolio'),
    ]))
    if contact_line:
        doc.add_paragraph().add_run(contact_line).font.size = Pt(9.5)

    doc.add_paragraph().add_run(datetime.now().strftime('%B %d, %Y')).font.size = Pt(10)

    if letter.get('company_name'):
        doc.add_paragraph().add_run(letter['company_name']).font.size = Pt(10)

    doc.add_paragraph()

    doc.add_paragraph(letter.get('salutation') or 'Dear Hiring Manager,')

    for para in letter.get('body_paragraphs', []):
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph(letter.get('closing') or 'Sincerely,')
    doc.add_paragraph(letter.get('full_name', ''))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route('/generate-cover-letter')
@login_required
def generate_cover_letter():
    return render_template('generate_cover_letter.html')


@app.route('/generate-cover-letter/build', methods=['POST'])
@login_required
def generate_cover_letter_build():
    if not openrouter_client:
        return jsonify({"error": "OPENROUTER_API_KEY is not configured on the server."}), 503

    input_data = request.get_json(silent=True) or {}
    if not (input_data.get('full_name') and input_data.get('target_role') and input_data.get('company_name')):
        return jsonify({"error": "Full name, target role, and company name are required."}), 400

    try:
        letter = _generate_cover_letter_content(input_data)
    except Exception as e:
        return jsonify({"error": f"Failed to generate cover letter: {e}"}), 502

    return jsonify({"letter": letter})


@app.route('/generate-cover-letter/render', methods=['POST'])
@login_required
def generate_cover_letter_render():
    data = request.get_json(silent=True) or {}
    return _render_cover_letter_html(data.get('template', 'modern'), data.get('letter') or {})


@app.route('/generate-cover-letter/download/pdf', methods=['POST'])
@login_required
def generate_cover_letter_download_pdf():
    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    letter = data.get('letter') or {}
    try:
        pdf_bytes = _html_to_pdf_bytes(_render_cover_letter_html(template, letter))
    except Exception as e:
        return jsonify({"error": f"Failed to build PDF: {e}"}), 500
    return send_file(
        io.BytesIO(pdf_bytes), mimetype='application/pdf',
        as_attachment=True, download_name=f"{_slugify(letter.get('full_name'))}-cover-letter.pdf",
    )


@app.route('/generate-cover-letter/download/docx', methods=['POST'])
@login_required
def generate_cover_letter_download_docx():
    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    letter = data.get('letter') or {}
    try:
        docx_bytes = _build_cover_letter_docx(template, letter)
    except Exception as e:
        return jsonify({"error": f"Failed to build Word document: {e}"}), 500
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name=f"{_slugify(letter.get('full_name'))}-cover-letter.docx",
    )


@app.route('/generate-cover-letter/save', methods=['POST'])
@login_required
def generate_cover_letter_save():
    if db is None:
        return jsonify({"error": "Database is not configured on the server."}), 503

    data = request.get_json(silent=True) or {}
    template = data.get('template', 'modern')
    input_data = data.get('input_data') or {}
    letter = data.get('letter') or {}

    if not letter:
        return jsonify({"error": "No generated cover letter to save."}), 400

    doc_ref = db.collection('cover_letters').document()
    doc_ref.set({
        'user_id': session['user_id'],
        'template': template,
        'full_name': letter.get('full_name', ''),
        'target_role': letter.get('target_role', ''),
        'company_name': letter.get('company_name', ''),
        'input_data': input_data,
        'generated_content': letter,
        'created_at': firestore.SERVER_TIMESTAMP,
    })

    return jsonify({"id": doc_ref.id})


@app.route('/cover-letter/<letter_id>/download/pdf')
@login_required
def cover_letter_download_pdf(letter_id):
    row = _get_owned_doc('cover_letters', letter_id)
    pdf_bytes = _html_to_pdf_bytes(_render_cover_letter_html(row['template'], row['generated_content']))
    return send_file(
        io.BytesIO(pdf_bytes), mimetype='application/pdf',
        as_attachment=True, download_name=f"{_slugify(row['full_name'])}-cover-letter.pdf",
    )


@app.route('/cover-letter/<letter_id>/download/docx')
@login_required
def cover_letter_download_docx(letter_id):
    row = _get_owned_doc('cover_letters', letter_id)
    docx_bytes = _build_cover_letter_docx(row['template'], row['generated_content'])
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name=f"{_slugify(row['full_name'])}-cover-letter.docx",
    )


@app.route('/cover-letter/<letter_id>/delete', methods=['POST'])
@login_required
def cover_letter_delete(letter_id):
    _get_owned_doc('cover_letters', letter_id)
    db.collection('cover_letters').document(letter_id).delete()
    return jsonify({"ok": True})


# --- PLAGIARISM CHECKER (PUBLIC, NO LOGIN REQUIRED) ---

MAX_PLAGIARISM_TEXT_CHARS = 8000


class PlagiarismMatch(BaseModel):
    excerpt: str = ""
    similarity_percentage: int = 0
    source_description: str = ""
    source_url: str = ""


class PlagiarismReport(BaseModel):
    originality_percentage: int = 100
    plagiarism_percentage: int = 0
    summary: str = ""
    matches: list[PlagiarismMatch] = []


def _extract_text_from_upload(file_storage):
    filename = (file_storage.filename or '').lower()
    raw = file_storage.read()

    if filename.endswith('.txt'):
        return raw.decode('utf-8', errors='ignore')

    if filename.endswith('.docx'):
        doc = Document(io.BytesIO(raw))
        return '\n'.join(p.text for p in doc.paragraphs)

    if filename.endswith('.pdf'):
        reader = PdfReader(io.BytesIO(raw))
        return '\n'.join((page.extract_text() or '') for page in reader.pages)

    raise ValueError('Unsupported file type. Please upload a .txt, .docx, or .pdf file.')


def _check_plagiarism_with_gemini(text, exclude_citations):
    citation_instruction = (
        "Ignore text that appears to be a direct quotation or citation (e.g. in "
        "quotation marks, or clearly attributed to another source) when judging "
        "originality - do not flag properly quoted/cited material as plagiarism."
        if exclude_citations else
        "Treat all text as the author's own claimed work, including quoted passages."
    )

    prompt = f"""You are a plagiarism detection assistant with access to Google Search.
Analyze the following text for originality. Search the web to check whether
passages closely match existing published content (articles, papers, websites).
{citation_instruction}

Break the text into meaningful excerpts/sentences and check each one for close
matches elsewhere on the web. For every excerpt with a notable match, report the
excerpt, an estimated similarity percentage (0-100), a short description of the
matching source, and its URL if you found one.

Then give:
- "originality_percentage": overall estimated originality (0-100, where 100 is
  fully original)
- "plagiarism_percentage": 100 minus originality, roughly weighted by how much
  of the text is covered by matches
- "summary": a 2-3 sentence plain-English summary of the findings
- "matches": array of the flagged excerpts as described above (empty array if
  no notable matches found)

Text to analyze:
\"\"\"
{text}
\"\"\"

Respond with ONLY a JSON object (no markdown, no commentary) with exactly these
keys: "originality_percentage", "plagiarism_percentage", "summary", "matches"
(each match an object with "excerpt", "similarity_percentage",
"source_description", "source_url")."""

    if genai_client:
        try:
            response = genai_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    tools=[types.Tool(google_search=types.GoogleSearch())],
                ),
            )
            return _extract_json_object(response.text), True
        except Exception:
            # Live grounded search unavailable (quota/billing) - fall through
            # to a non-grounded assessment on OpenRouter instead.
            pass

    fallback_prompt = prompt.replace(
        "You are a plagiarism detection assistant with access to Google Search.",
        "You are a plagiarism detection assistant. Live web search is "
        "unavailable, so use your general knowledge to flag text that "
        "resembles widely known published content instead of live results.",
    )
    fallback_prompt += (
        '\n\nRespond with ONLY a JSON object (no markdown, no commentary) with '
        'exactly these keys: "originality_percentage" (int), '
        '"plagiarism_percentage" (int), "summary" (string), "matches" (array of '
        'objects with "excerpt", "similarity_percentage" (int), '
        '"source_description", "source_url").'
    )
    raw = _openrouter_generate(fallback_prompt, json_mode=True)
    data = _extract_json_object(raw)
    return PlagiarismReport(**data).model_dump(), False


@app.route('/plagiarism-checker')
def plagiarism_checker():
    return render_template('plagiarism_checker.html')


@app.route('/plagiarism-checker/check', methods=['POST'])
def plagiarism_checker_check():
    if not genai_client and not openrouter_client:
        return jsonify({"error": "Neither GEMINI_API_KEY nor OPENROUTER_API_KEY is configured on the server."}), 503

    if request.content_type and 'multipart/form-data' in request.content_type:
        exclude_citations = request.form.get('exclude_citations') == 'true'
        uploaded = request.files.get('file')
        if uploaded and uploaded.filename:
            try:
                text = _extract_text_from_upload(uploaded)
            except ValueError as e:
                return jsonify({"error": str(e)}), 400
        else:
            text = request.form.get('text') or ''
    else:
        data = request.get_json(silent=True) or {}
        text = data.get('text') or ''
        exclude_citations = bool(data.get('exclude_citations'))

    text = text.strip()
    if not text:
        return jsonify({"error": "Please paste some text or upload a document."}), 400
    if len(text) < 50:
        return jsonify({"error": "Please provide at least a few sentences to check."}), 400

    truncated = len(text) > MAX_PLAGIARISM_TEXT_CHARS
    text = text[:MAX_PLAGIARISM_TEXT_CHARS]

    try:
        report, live = _check_plagiarism_with_gemini(text, exclude_citations)
    except Exception as e:
        return jsonify({"error": f"Failed to check plagiarism: {e}"}), 502

    return jsonify({"report": report, "live_search": live, "truncated": truncated, "text": text})


@app.route('/plagiarism-checker/download/pdf', methods=['POST'])
def plagiarism_checker_download_pdf():
    data = request.get_json(silent=True) or {}
    text = data.get('text') or ''
    report = data.get('report') or {}

    if not report:
        return jsonify({"error": "No report to export."}), 400

    html = render_template(
        'plagiarism_report_pdf.html', report=report, text=text,
        generated_on=datetime.now().strftime('%B %d, %Y'),
    )
    try:
        pdf_bytes = _html_to_pdf_bytes(html)
    except Exception as e:
        return jsonify({"error": f"Failed to build PDF: {e}"}), 500

    return send_file(
        io.BytesIO(pdf_bytes), mimetype='application/pdf',
        as_attachment=True, download_name='plagiarism-report.pdf',
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug_mode = os.environ.get("FLASK_DEBUG", "true").lower() == "true"
    app.run(host="0.0.0.0", port=port, debug=debug_mode)


# --- ADMIN PANEL -----------------------------------------------------------
# Unlinked by design: nothing in the public navigation points at /admin, so it
# is reached by typing the URL. That is obscurity, not access control - the
# admin_required decorator is what actually guards these pages.

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        if db is None:
            flash('Database is not configured on the server.', 'error')
            return render_template('admin/login.html', next='')

        email = request.form.get('email', '')
        password = request.form.get('password', '')
        if adminstore.check_admin_login(db, email, password):
            session['admin_email'] = email.strip().lower()
            return redirect(request.form.get('next') or url_for('admin_dashboard'))
        flash('Invalid admin credentials.', 'error')

    return render_template('admin/login.html', next=request.args.get('next', ''))


@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_email', None)
    return redirect(url_for('admin_login'))


@app.route('/admin')
@admin_required
def admin_dashboard():
    return render_template('admin/dashboard.html',
                           analytics=adminstore.build_analytics(db))


@app.route('/admin/jobs', methods=['GET', 'POST'])
@admin_required
def admin_jobs():
    if request.method == 'POST':
        form = {k: (v or '').strip() for k, v in request.form.items()}
        if not form.get('title') or not form.get('company') or not form.get('apply_url'):
            flash('Title, company and application link are all required.', 'error')
        else:
            post_id = form.get('post_id')
            if post_id:
                adminstore.update_job_post(db, post_id, form)
                flash('Job post updated.', 'success')
            else:
                adminstore.create_job_post(db, form, session['admin_email'])
                flash('Job posted — it is live on the home page.', 'success')
            return redirect(url_for('admin_jobs'))

    return render_template(
        'admin/jobs.html',
        posts=adminstore.list_job_posts(db),
        categories=adminstore.JOB_CATEGORIES,
        employment_types=adminstore.EMPLOYMENT_TYPES,
        company_types=adminstore.COMPANY_TYPES,
        editing=adminstore.get_job_post(db, request.args.get('edit')) if request.args.get('edit') else None,
    )


@app.route('/admin/jobs/<post_id>/delete', methods=['POST'])
@admin_required
def admin_jobs_delete(post_id):
    adminstore.delete_job_post(db, post_id)
    flash('Job post deleted.', 'success')
    return redirect(url_for('admin_jobs'))


@app.route('/admin/users')
@admin_required
def admin_users():
    return render_template('admin/users.html', users=adminstore.list_users(db))


@app.route('/admin/users/<email>/reset', methods=['POST'])
@admin_required
def admin_users_reset(email):
    new_password = request.form.get('new_password', '')
    if len(new_password) < 6:
        flash('That password is too short — use at least 6 characters.', 'error')
    else:
        adminstore.reset_user_password(db, email, new_password)
        flash(f'Password reset for {email}. Pass it to them over a channel they already trust.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/users/<email>/email', methods=['POST'])
@admin_required
def admin_users_email(email):
    new_email = request.form.get('new_email', '').strip().lower()
    if not new_email or '@' not in new_email:
        flash('Enter a valid email address.', 'error')
    else:
        try:
            adminstore.update_user_email(db, email, new_email)
            flash(f'{email} is now {new_email}.', 'success')
        except ValueError as e:
            flash(str(e), 'error')
    return redirect(url_for('admin_users'))


@app.route('/admin/analytics')
@admin_required
def admin_analytics():
    return render_template('admin/analytics.html',
                           analytics=adminstore.build_analytics(db))


@app.route('/admin/announcements', methods=['GET', 'POST'])
@admin_required
def admin_announcements():
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        body = request.form.get('body', '').strip()
        if not title or not body:
            flash('An announcement needs both a title and a message.', 'error')
        else:
            adminstore.create_announcement(db, title, body, request.form.get('level', 'info'))
            flash('Announcement published — signed-in users see it on their next page load.', 'success')
            return redirect(url_for('admin_announcements'))

    return render_template('admin/announcements.html',
                           announcements=adminstore.list_announcements(db))


@app.route('/admin/announcements/<ann_id>/delete', methods=['POST'])
@admin_required
def admin_announcements_delete(ann_id):
    adminstore.delete_announcement(db, ann_id)
    flash('Announcement removed.', 'success')
    return redirect(url_for('admin_announcements'))


@app.route('/admin/settings', methods=['GET', 'POST'])
@admin_required
def admin_settings():
    if request.method == 'POST':
        adminstore.save_settings(
            db,
            request.form.get('app_name', ''),
            request.form.get('footer_signature', ''),
        )
        flash('Settings saved — they apply across the site immediately.', 'success')
        return redirect(url_for('admin_settings'))

    return render_template('admin/settings.html', settings=adminstore.get_settings(db))


@app.route('/admin/profile', methods=['GET', 'POST'])
@admin_required
def admin_profile():
    if request.method == 'POST':
        current = request.form.get('current_password', '')
        new_email = request.form.get('email', '').strip().lower()
        new_password = request.form.get('new_password', '')

        # Re-check the current password before either change: an unattended
        # session should not be enough to take the account over.
        if not adminstore.check_admin_login(db, session['admin_email'], current):
            flash('Current password is incorrect.', 'error')
        elif new_password and len(new_password) < 6:
            flash('New password must be at least 6 characters.', 'error')
        elif new_email and '@' not in new_email:
            flash('Enter a valid email address.', 'error')
        else:
            adminstore.update_admin(db, new_email or None, new_password or None)
            if new_email:
                session['admin_email'] = new_email
            flash('Admin account updated.', 'success')
            return redirect(url_for('admin_profile'))

    return render_template('admin/profile.html', admin=adminstore.get_admin(db))


@app.route('/announcements/seen', methods=['POST'])
@login_required
def announcements_seen():
    """Marks the bell as read. Stored on the user so it survives a new login."""
    stamp = datetime.now(timezone.utc).timestamp()
    session['ann_seen_at'] = stamp
    if db is not None:
        try:
            db.collection('users').document(session['user_id']).update(
                {'ann_seen_at': stamp})
        except Exception:
            pass   # session already updated; persistence is best-effort
    return jsonify({"ok": True})


@app.route('/admin/questions', methods=['GET', 'POST'])
@admin_required
def admin_questions():
    """
    The questionnaire module: add practice questions per company, and report
    on the bank as a whole.

    Companies are keyed by a slug derived from the typed name, so entering a
    company that already exists stacks onto it rather than creating a second
    one - which is what makes "add to their stack" work without a picker.
    """
    if request.method == 'POST':
        form = {k: (v or '').strip() for k, v in request.form.items()}
        options = [o.strip() for o in request.form.getlist('options') if o.strip()]

        if not form.get('company_name') or not form.get('prompt'):
            flash('Company name and the question itself are both required.', 'error')
        elif form.get('type') == 'MCQ' and len(options) < 2:
            flash('An MCQ needs at least two options.', 'error')
        elif form.get('type') == 'MCQ' and not form.get('answer'):
            flash('Choose which option is correct.', 'error')
        else:
            prepbank.add_question(db, {**form, 'options': options}, session['admin_email'])
            slug = prepbank.slugify(form['company_name'])
            existing = len(prepbank.custom_for_slug(db, slug)) + len(prepbank.bank_questions(slug))
            flash(f"Question added to {form['company_name']} — {existing} now in their stack.",
                  'success')
            return redirect(url_for('admin_questions'))

    return render_template(
        'admin/questions.html',
        analytics=prepbank.build_analytics(db),
        question_types=prepbank.QUESTION_TYPES,
        difficulties=prepbank.DIFFICULTIES,
        known_companies=sorted(prepbank.company_names(db).values()),
    )


@app.route('/admin/questions/<question_id>/delete', methods=['POST'])
@admin_required
def admin_questions_delete(question_id):
    prepbank.delete_question(db, question_id)
    flash('Question removed.', 'success')
    return redirect(url_for('admin_questions'))


@app.route('/admin/questions/template.xlsx')
@admin_required
def admin_questions_template():
    """The blank questionnaire sheet, generated rather than stored on disk."""
    buffer = io.BytesIO()
    prepbank.build_template_workbook().save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='questionnaire-template.xlsx',
    )


@app.route('/admin/questions/upload', methods=['POST'])
@admin_required
def admin_questions_upload():
    upload = request.files.get('sheet')
    if not upload or not upload.filename:
        flash('Choose a spreadsheet to upload.', 'error')
        return redirect(url_for('admin_questions'))
    if not upload.filename.lower().endswith(('.xlsx', '.xlsm')):
        flash('That is not an Excel file — upload the .xlsx template.', 'error')
        return redirect(url_for('admin_questions'))

    try:
        rows, errors = prepbank.parse_workbook(upload.stream)
    except Exception as e:
        flash(f'Could not read that file: {e}', 'error')
        return redirect(url_for('admin_questions'))

    written = prepbank.add_many(db, rows, session['admin_email']) if rows else 0

    if written:
        flash(f'Imported {written} question{"" if written == 1 else "s"}.', 'success')
    if errors:
        # Report every rejected row rather than a count, so the admin can fix
        # the sheet instead of guessing which line was wrong.
        shown = errors[:10]
        more = f' (+{len(errors) - 10} more)' if len(errors) > 10 else ''
        flash('Skipped: ' + ' '.join(shown) + more, 'error')
    if not written and not errors:
        flash('That sheet had no rows to import.', 'error')

    return redirect(url_for('admin_questions'))


@app.errorhandler(404)
def page_not_found(_e):
    """
    Any unknown route, plus every abort(404) already in the app.

    The status code is preserved deliberately - a pretty page returned as 200
    tells crawlers and monitoring the URL is fine, which is how dead links get
    indexed.
    """
    return render_template('404.html'), 404
